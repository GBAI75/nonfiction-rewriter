import streamlit as st
import pandas as pd
from openai import OpenAI
from openai import RateLimitError, APIError
from io import BytesIO
from docx import Document
from docx.shared import Pt

# Set page config
st.set_page_config(page_title="Non-Fiction Rewriter", layout="centered")

# Initialize OpenAI client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

st.title("📘 Non-Fiction Paragraph Rewriter")

# Session state to store saved paragraphs
if "entries" not in st.session_state:
    st.session_state.entries = []

# Prompt choices
st.subheader("Step 1: Choose Rewrite Style")
prompt_option = st.radio(
    "Select how you'd like the assistant to rewrite your text:",
    [
        "1) Interpret meaning + expand using online patterns (Default)",
        "2) Improve coherence and grammar",
        "3) Custom prompt"
    ]
)

base_prompt = (
    "Improve the text by first interpreting its meaning, extrapolating its fuller intent, then expanding the depth of thought and argument through similar thoughts found in an online search you can run and then drafting a section of a chapter in a book about listening to other perspectives and to nuances in our search toward common grounds to accelerate solutions to the climate crisis. Draft as if you are the author of the book. Ensure there is no plagiarism in the exact text used if you find arguments and statement online. Ok to use quotes if they are properly referenced to original author and great to quote facts but put footnotes to references."
)

if prompt_option.startswith("3"):
    custom_prompt = st.text_area("Enter your custom prompt", value=base_prompt, height=200)
else:
    custom_prompt = (
        base_prompt if prompt_option.startswith("1") else
        "Improve the text by redeveloping it for more coherent thought flow and impeccable style and grammar."
    )

# Input text box or voice input
st.subheader("Step 2: Input Your Draft")
input_mode = st.radio("Choose input method:", ["Type", "Dictate (speech-to-text)"])

if input_mode == "Type":
    input_text = st.text_area("Write your messy draft here", height=200)
else:
    input_text = st.text_area("Speak and paste your dictation here (voice input capture not supported directly in Streamlit yet)", height=200)

# Rewrite logic
if st.button("✍️ Rewrite into polished paragraph"):
    if input_text.strip() == "":
        st.warning("Please enter some text.")
    else:
        with st.spinner("Rewriting..."):
            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": custom_prompt},
                        {"role": "user", "content": input_text}
                    ],
                    temperature=0.7
                )
                rewritten = response.choices[0].message.content
                st.session_state.latest = rewritten
                st.success("Here's your rewritten paragraph:")
                st.write(rewritten)
            except RateLimitError:
                st.error("🚫 Rate limit reached. Please wait a moment and try again.")
            except APIError as e:
                st.error(f"⚠️ OpenAI API error: {str(e)}")
            except Exception as e:
                st.error(f"❌ Unexpected error: {str(e)}")

# Save title, keywords
if "latest" in st.session_state:
    with st.form("save_form"):
        title = st.text_input("Give this paragraph a title")
        keywords = st.text_input("Add some keywords (comma-separated)")
        submitted = st.form_submit_button("💾 Save this paragraph")
        if submitted:
            st.session_state.entries.append({
                "title": title,
                "keywords": keywords,
                "paragraph": st.session_state.latest
            })
            st.success("Saved! You can add more or download all below.")
            del st.session_state.latest

# Display saved entries
if st.session_state.entries:
    st.subheader("📚 Saved Paragraphs")
    for i, entry in enumerate(st.session_state.entries):
        st.markdown(f"**{i+1}. {entry['title']}**  \n*Keywords:* {entry['keywords']}  \n\n{entry['paragraph']}")

    # Word document export
    def generate_docx(entries):
        doc = Document()
        for entry in entries:
            doc.add_heading(entry['title'], level=1)
            p = doc.add_paragraph()
            run = p.add_run("Keywords: ")
            run.bold = True
            run.font.size = Pt(10)
            run2 = p.add_run(entry['keywords'])
            run2.font.size = Pt(10)
            doc.add_paragraph(entry['paragraph'])
            doc.add_paragraph("\n")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    docx_buffer = generate_docx(st.session_state.entries)
    st.download_button(
        "📄 Download All as Word Document",
        data=docx_buffer,
        file_name="rewritten_paragraphs.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

